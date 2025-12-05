from dotenv import load_dotenv
load_dotenv()
from flask import Flask, jsonify, request
import os
import json
import glob
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from pathlib import Path
from datetime import datetime, timezone
import logging
import time
import fitz  # PyMuPDF
import pikepdf # For writing PDF metadata
import uuid    # For generating GUIDs
import base64  # To encode images for API call
from openai import OpenAI # To call the vision model
from docling.document_converter import DocumentConverter, PdfFormatOption, WordFormatOption
from docling.datamodel.pipeline_options import (
    PdfPipelineOptions, TableStructureOptions, TableFormerMode
)
from docling.datamodel.base_models import InputFormat
import pdfplumber
import tempfile
import zipfile
from chonkie import SemanticChunker # type: ignore
import hashlib
from io import BytesIO
try:
    from PIL import Image  # for thumbnails
    PIL_AVAILABLE = True
except Exception:
    PIL_AVAILABLE = False


# --- Basic Configuration ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("flask_app.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)

BASE_DIR = "/mnt/c/Users/Anderson/Documents/n8n/kmc-rag/test_docs"
print(f"BASE_DIR is set to: {BASE_DIR}")
JSON_CACHE_DIR = os.path.join(BASE_DIR, "json_cache")
IMAGE_OUTPUT_DIR = os.path.join(BASE_DIR, "image_output")
os.makedirs(JSON_CACHE_DIR, exist_ok=True)

SUPPORTED_EXTENSIONS = ['.docx', '.pdf']

# --- FileID Registry (Global Cache) ---
# Maps fileID -> file_path for quick lookups
FILE_ID_REGISTRY = {}
FILE_PATH_CACHE = {}  # Maps file_path -> fileID for reverse lookups

# --- ALL METADATA FIELDS ALWAYS INCLUDED ---
ALL_METADATA_FIELDS = [
    "file_id",
    "filename",
    "full_path",
    "subfolder",
    "file_extension",
    "file_size",
    "file_size_mb",
    "word_count",
    "character_count",
    "paragraph_count",
    "table_count",
    "lastModified",
    "createdDate",
    "has_guid",
    "guid_source",
    "DateReviewed",
    "DateNext",
    "ai_title",
    "ai_description"
]

def register_file_id(file_id: str, file_path: str) -> None:
    """Register a fileID to file_path mapping."""
    FILE_ID_REGISTRY[file_id] = file_path
    FILE_PATH_CACHE[file_path] = file_id
    logger.debug(f"Registered fileID: {file_id} -> {file_path}")

def resolve_file_path(file_id: str) -> str | None:
    """Resolve a fileID to its file_path. Returns None if not found."""
    return FILE_ID_REGISTRY.get(file_id)

def get_file_id_or_register(file_path: str) -> str:
    """Get fileID for a path, or register it if not already present."""
    if file_path in FILE_PATH_CACHE:
        return FILE_PATH_CACHE[file_path]
    
    ext = os.path.splitext(file_path)[1].lower()
    guid = None
    
    try:
        if ext == '.pdf':
            guid = read_pdf_guid(file_path)
        elif ext == '.docx':
            guid = extract_docx_guid(file_path)
    except Exception as e:
        logger.warning(f"Failed to extract GUID from {file_path}: {e}")
    
    if not guid:
        guid = os.path.splitext(os.path.basename(file_path))[0].replace(' ', '_').replace('-', '_').lower()
    
    register_file_id(guid, file_path)
    return guid

def get_image_descriptions_from_api(image_info_list):
    # No-op stub: image analysis handled by /docling/convert-file two-step flow.
    return []

# --- Error Handling Helper ---

def create_error_response(file_id_or_path, error_message, error_type="CONVERSION_ERROR", details=None):
    """
    Creates a standardized error response that returns 200 OK with error flag.
    Accepts either fileID or file_path as first parameter.
    
    Args:
        file_id_or_path: FileID or file path that failed
        error_message: Human-readable error message
        error_type: Type of error (CONVERSION_ERROR, PARSING_ERROR, FILE_NOT_FOUND, etc.)
        details: Optional dict with additional error details
    
    Returns:
        tuple: (jsonify response dict, 200 status code)
    """
    # Determine if input is fileID or path
    if file_id_or_path and file_id_or_path.startswith('/'):
        file_path = file_id_or_path
        file_id = FILE_PATH_CACHE.get(file_path)
        if not file_id:
            file_id = get_file_id_or_register(file_path) if os.path.exists(file_path) else None
    else:
        file_id = file_id_or_path
        file_path = resolve_file_path(file_id)
    
    # Attempt to include file timestamps (UTC) when available
    last_modified = None
    created_date = None
    try:
        if file_path and os.path.exists(file_path):
            st = os.stat(file_path)
            last_modified = datetime.fromtimestamp(st.st_mtime, tz=timezone.utc).isoformat()
            created_date = datetime.fromtimestamp(st.st_ctime, tz=timezone.utc).isoformat()
    except Exception as e:
        logger.warning(f"Failed to stat file for error response timestamps: {e}")

    response = {
        "success": False,
        "error": True,
        "error_type": error_type,
        "error_message": error_message,
        "file_id": file_id,
        "file_path": file_path,
        "filename": os.path.basename(file_path) if file_path else "unknown",
        "conversion_status": "failed",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "lastModified": last_modified,
        "createdDate": created_date
    }
    if details:
        response["error_details"] = details
    
    logger.error(f"[{error_type}] {file_id}: {error_message}")
    return jsonify(response), 200

# --- Helper Functions ---

def make_complete_metadata(file_path, extra_data=None):
    """Create complete metadata dictionary with all fields, using None for missing values."""
    file_stats = os.stat(file_path)
    parent_dir = os.path.basename(os.path.dirname(file_path))
    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()

    # Base metadata that's always available
    metadata = {
        "filename": filename,
        "full_path": file_path,
        "subfolder": parent_dir,
        "file_extension": ext,
        "file_size": file_stats.st_size,
        "file_size_mb": round(file_stats.st_size / (1024 * 1024), 2),
        "lastModified": datetime.fromtimestamp(file_stats.st_mtime, tz=timezone.utc).isoformat(),
        "createdDate": datetime.fromtimestamp(file_stats.st_ctime, tz=timezone.utc).isoformat(),
    }

    # Initialize all fields from ALL_METADATA_FIELDS with None
    complete_metadata = {field: None for field in ALL_METADATA_FIELDS}
    
    # Update with base metadata
    complete_metadata.update(metadata)
    
    # Update with any extra data provided
    if extra_data and isinstance(extra_data, dict):
        complete_metadata.update(extra_data)

    return complete_metadata

def extract_docx_guid(docx_path):
    try:
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            settings_path = 'word/settings.xml'
            if settings_path not in docx_zip.namelist():
                return None
            with docx_zip.open(settings_path) as settings_file:
                root = ET.fromstring(settings_file.read())
            namespaces = {
                'w15': 'http://schemas.microsoft.com/office/word/2012/wordml'
            }
            doc_id_element = root.find('.//w15:docId', namespaces)
            if doc_id_element is not None:
                doc_id = doc_id_element.get('{http://schemas.microsoft.com/office/word/2012/wordml}val')
                if doc_id:
                    return doc_id.strip('{}')
        return None
    except Exception as e:
        logger.warning(f"Could not extract GUID from {docx_path}: {e}")
        return None

def get_or_create_pdf_guid(pdf_path: str) -> str:
    """Reads a custom GUID from a PDF's metadata. If not present, generates one and saves it."""
    guid_key = '/AppGUID'
    try:
        with pikepdf.open(pdf_path, allow_overwriting_input=True) as pdf:
            docinfo = pdf.docinfo
            if guid_key in docinfo:
                logger.info(f"Found existing GUID in {pdf_path}")
                return str(docinfo[guid_key])

            logger.info(f"No GUID found in {pdf_path}. Generating a new one.")
            new_guid = str(uuid.uuid4())
            docinfo[guid_key] = new_guid
            pdf.save()
            logger.info(f"Saved new GUID to {pdf_path}")
            return new_guid
            
    except Exception as e:
        logger.error(f"Pikepdf failed to read or write GUID for {pdf_path}: {e}", exc_info=True)
        filename_without_ext = os.path.splitext(os.path.basename(pdf_path))[0]
        return filename_without_ext.replace(' ', '_').replace('-', '_').lower()

def read_pdf_guid(pdf_path: str) -> str | None:
    """Reads a custom GUID from a PDF's metadata without modifying the file."""
    guid_key = '/AppGUID'
    try:
        with pikepdf.open(pdf_path) as pdf:
            if guid_key in pdf.docinfo:
                return str(pdf.docinfo[guid_key])
        return None
    except Exception as e:
        logger.warning(f"Could not read GUID from {pdf_path} with pikepdf: {e}")
        return None

def extract_custom_metadata(doc_path: str):
    date_reviewed, date_next = None, None
    try:
        with zipfile.ZipFile(doc_path, 'r') as z:
            if 'docProps/custom.xml' not in z.namelist():
                return date_reviewed, date_next
            xml_content = z.read('docProps/custom.xml')
            root = ET.fromstring(xml_content)
            
            CP_NAMESPACE = 'http://schemas.openxmlformats.org/officeDocument/2006/custom-properties'
            VT_NAMESPACE = 'http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes'
            FQN_PROPERTY = f'{{{CP_NAMESPACE}}}property'
            FQN_FILETIME = f'{{{VT_NAMESPACE}}}filetime'

            for prop in root.findall(FQN_PROPERTY):
                name = prop.attrib.get('name')
                if name == 'DateReviewed':
                    value_elem = prop.find(FQN_FILETIME)
                    if value_elem is not None and value_elem.text:
                        date_reviewed = datetime.fromisoformat(value_elem.text.replace('Z', '+00:00')).isoformat()
                elif name == 'DateNext':
                    value_elem = prop.find(FQN_FILETIME)
                    if value_elem is not None and value_elem.text:
                        date_next = datetime.fromisoformat(value_elem.text.replace('Z', '+00:00')).isoformat()
    except Exception as e:
        logger.warning(f"Failed to extract custom metadata from {doc_path}: {e}")
    return date_reviewed, date_next

def is_docx_complex(docx_path,
                    min_tables_for_complex=1,
                    min_images_for_complex=1,
                    max_paragraphs_for_simple=300,
                    avg_run_threshold=2.5):
    """Returns True if the DOCX should be processed by Docling (complex)."""
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            namelist = z.namelist()
            media_files = [n for n in namelist if n.startswith('word/media/')]
            embedded_objects = [n for n in namelist if n.startswith('word/embeddings') or n.endswith('.bin')]
            has_images = len(media_files) >= min_images_for_complex
            has_embedded = len(embedded_objects) > 0

        doc = Document(docx_path)
        num_tables = len(doc.tables)
        num_paragraphs = len(doc.paragraphs)

        total_runs = 0
        paragraphs_sampled = 0
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            paragraphs_sampled += 1
            try:
                total_runs += len(p.runs)
            except Exception:
                total_runs += 1
            if paragraphs_sampled >= 200:
                break
        avg_runs = (total_runs / paragraphs_sampled) if paragraphs_sampled else 1

        if num_tables >= min_tables_for_complex:
            return True
        if has_images or has_embedded:
            return True
        if num_paragraphs > max_paragraphs_for_simple:
            return True
        if avg_runs > avg_run_threshold:
            return True

        return False
    except Exception as e:
        logger.warning(f"is_docx_complex failed for {docx_path}: {e}", exc_info=True)
        return True

def extract_tables_with_pdfplumber(pdf_path):
    """Extract tables from PDF using pdfplumber"""
    tables = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_tables = page.extract_tables()
                for table in page_tables or []:
                    if table and len(table) > 0:
                        clean_table = []
                        for row in table:
                            clean_row = [cell.strip() if cell else "" for cell in row]
                            if any(clean_row):
                                clean_table.append(clean_row)
                        
                        if clean_table:
                            tables.append({
                                "page": page_num + 1,
                                "headers": clean_table[0] if clean_table else [],
                                "rows": clean_table[1:] if len(clean_table) > 1 else [],
                                "raw_data": clean_table
                            })
    except Exception as e:
        logger.warning(f"pdfplumber table extraction failed for {pdf_path}: {e}")
    return tables

def format_table_for_rag(table_data):
    """Format table data for better RAG retrieval"""
    if not table_data or not table_data.get("raw_data"):
        return ""
    
    raw_data = table_data["raw_data"]
    formatted_lines = []
    
    formatted_lines.append("\n--- TABLE START ---")
    
    for row_index, row in enumerate(raw_data):
        if row and any(cell.strip() for cell in row if cell):
            row_text = " | ".join([cell for cell in row if cell and cell.strip()])
            formatted_lines.append(row_text)
            
            if row_index == 0:
                formatted_lines.append("-" * 40)
    
    formatted_lines.append("--- TABLE END ---\n")
    return "\n".join(formatted_lines)

def generate_ai_title_description(full_text: str, fallback_filename: str) -> tuple[str, str]:
    def sample(text: str, max_chars: int = 12000) -> str:
        if len(text) <= max_chars:
            return text
        head = text[: max_chars // 2]
        tail = text[-max_chars // 2 :]
        return head + "\n...\n" + tail

    try:
        client = OpenAI()
    except Exception as e:
        logger.warning(f"OpenAI client unavailable, will fallback: {e}")
        client = None

    default_title = Path(fallback_filename).stem.replace("_", " ").replace("-", " ").strip()
    default_desc = "No AI description available."

    if not client:
        return default_title, default_desc

    prompt = (
        "You are given extracted document text. "
        "Return a concise JSON with fields 'title' and 'description'. "
        "Title <= 12 words, description 1–2 sentences. "
        "Do not include code fences.\n\n"
        f"TEXT START\n{sample(full_text)}\nTEXT END"
    )

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=200,
        )
        content = resp.choices[0].message.content.strip()
        data = json.loads(content) if content.startswith("{") else {}
        title = (data.get("title") or default_title).strip()
        desc = (data.get("description") or default_desc).strip()
        return title, desc
    except Exception as e:
        logger.warning(f"AI title/description generation failed, using fallback: {e}")
        return default_title, default_desc

# def get_image_descriptions_from_api(image_info_list: list) -> list:
#     """Takes a list of image paths and calls a vision model for each."""
#     try:
#         client = OpenAI()
#     except Exception as e:
#         logger.error(f"OpenAI client failed to initialize. Is OPENAI_API_KEY set? Error: {e}")
#         return []

#     descriptions = []
#     for image_info in image_info_list:
#         path = image_info.get("path")
#         try:
#             with open(path, "rb") as image_file:
#                 base64_image = base64.b64encode(image_file.read()).decode('utf-8')

#             response = client.chat.completions.create(
#                 model="gpt-4o",
#                 messages=[
#                     {"role": "user", "content": [
#                         {"type": "text", "text": "Describe this image from a document in detail. If it's a chart or graph, explain what it shows."},
#                         {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
#                     ]}
#                 ],
#                 max_tokens=300,
#             )
#             descriptions.append({"page": image_info.get("page"), "description": response.choices[0].message.content})
#         except Exception as e:
#             logger.error(f"Failed to get description for image {path}: {e}")
#     return descriptions

# --- Core Conversion Functions ---

def docx_to_json(docx_path):
    logger.info(f"Starting DOCX conversion for file: {docx_path}")
    try:
        doc = Document(docx_path)
        filename_without_ext, _ = os.path.splitext(os.path.basename(docx_path))
        filename_based_id = filename_without_ext.replace(' ', '_').replace('-', '_').lower()
        
        guid = extract_docx_guid(docx_path)
        file_id = guid or filename_based_id
        has_guid = guid is not None

        date_reviewed, date_next = extract_custom_metadata(docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Create complete metadata with all fields
        extra_data = {
            "file_id": file_id,
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "word_count": sum(len(p.split()) for p in paragraphs),
            "character_count": sum(len(p) for p in paragraphs),
            "has_guid": has_guid,
            "guid_source": "docx_settings" if has_guid else "filename_fallback",
            "DateReviewed": date_reviewed,
            "DateNext": date_next
        }
        
        metadata = make_complete_metadata(docx_path, extra_data)

        content = {
            "file_id": file_id,
            "filename": os.path.basename(docx_path),
            "full_path": docx_path,
            "subfolder": os.path.basename(os.path.dirname(docx_path)),
            "paragraphs": paragraphs,
            "tables": [[cell.text.strip() for cell in row.cells] for table in doc.tables for row in table.rows],
            "metadata": metadata
        }
        return content
    except Exception as e:
        logger.error(f"Failed to convert {docx_path}: {e}", exc_info=True)
        return {"error": f"Failed to convert {docx_path}: {str(e)}"}

def pdf_to_json(pdf_path):
    logger.info(f"Starting PDF conversion for file: {pdf_path}")
    try:
        file_id = get_or_create_pdf_guid(pdf_path)
        
        # Image extraction
        doc = fitz.open(pdf_path)
        image_info = []
        image_output_folder = os.path.join(IMAGE_OUTPUT_DIR, file_id)
        os.makedirs(image_output_folder, exist_ok=True)

        for page_index in range(len(doc)):
            page = doc[page_index]
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                image_filename = f"page{page_index+1}_img{img_index+1}.{image_ext}"
                image_save_path = os.path.join(image_output_folder, image_filename)
                
                with open(image_save_path, "wb") as img_file:
                    img_file.write(image_bytes)
                
                image_info.append({
                    "path": image_save_path,
                    "page": page_index + 1,
                    "xref": xref
                })

        # Extract tables using pdfplumber
        extracted_tables = extract_tables_with_pdfplumber(pdf_path)
        
        # Extract text using PyMuPDF
        full_text = "".join(page.get_textpage().extractText() for page in doc)
        paragraphs = [p.strip() for p in full_text.split('\n') if p.strip()]
        
        # Convert tables to format expected by chunking code
        tables_for_chunking = []
        for table in extracted_tables:
            if table.get("raw_data"):
                tables_for_chunking.extend(table["raw_data"])

        # Create complete metadata with all fields
        extra_data = {
            "file_id": file_id,
            "paragraph_count": len(paragraphs),
            "table_count": len(extracted_tables),
            "word_count": len(full_text.split()),
            "character_count": len(full_text),
            "has_guid": True,
            "guid_source": "pdf_metadata",
            "DateReviewed": None,
            "DateNext": None
        }
        
        metadata = make_complete_metadata(pdf_path, extra_data)
        
        content = {
            "file_id": file_id,
            "filename": os.path.basename(pdf_path),
            "full_path": pdf_path,
            "subfolder": os.path.basename(os.path.dirname(pdf_path)),
            "paragraphs": paragraphs,
            "tables": tables_for_chunking,
            "images": image_info,
            "extracted_tables": extracted_tables,
            "metadata": metadata
        }
        
        doc.close()
        return content
    except Exception as e:
        logger.error(f"Failed to convert {pdf_path}: {e}", exc_info=True)
        return {"error": f"Failed to convert {pdf_path}: {str(e)}"}

def convert_file_to_json(file_path):
    _, extension = os.path.splitext(file_path)
    extension = extension.lower()
    
    if extension == '.docx':
        return docx_to_json(file_path)
    elif extension == '.pdf':
        return pdf_to_json(file_path)
    else:
        logger.warning(f"Unsupported file type: {extension} for file {file_path}")
        return {"error": f"Unsupported file type: {extension}"}

# --- Caching and File System Functions ---

def find_all_supported_files():
    supported_files = []
    
    # Use os.walk for more reliable file discovery
    for root, dirs, files in os.walk(BASE_DIR):
        for filename in files:
            # Skip temporary files
            if filename.startswith('~$'):
                continue
                
            # Check if file has supported extension (case-insensitive)
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext in [ext.lower() for ext in SUPPORTED_EXTENSIONS]:
                file_path = os.path.join(root, filename)
                try:
                    file_stats = os.stat(file_path)
                    parent_dir = os.path.basename(os.path.dirname(file_path))
                    supported_files.append({
                        "filename": filename,
                        "full_path": file_path,
                        "subfolder": parent_dir,
                        "lastModified": datetime.fromtimestamp(file_stats.st_mtime, tz=timezone.utc).isoformat()
                    })
                except OSError as e:
                    logger.warning(f"Could not stat file {file_path}: {e}")
                    continue
    
    return supported_files
def get_cache_path(file_path):
    filename_without_ext = os.path.splitext(os.path.basename(file_path))[0]
    parent_dir = os.path.basename(os.path.dirname(file_path))
    cache_name = f"{parent_dir}_{filename_without_ext}.json".replace(' ', '_').replace('-', '_')
    return os.path.join(JSON_CACHE_DIR, cache_name)

def is_cache_valid(file_path, cache_path):
    if not os.path.exists(cache_path):
        return False
    return os.path.getmtime(cache_path) > os.path.getmtime(file_path)

def perform_chunking(input_text, complete_metadata, options):
    """Performs semantic chunking and enriches with COMPLETE metadata for RAG."""
    try:
        chunker = SemanticChunker(
            threshold=options.get("threshold", 0.75),
            chunk_size=options.get("chunk_size", 512),
            device_map="cpu"
        )
        raw_chunks = chunker.chunk(input_text)

        enriched_chunks = []
        for chunk in raw_chunks:
            # Each chunk gets ALL metadata fields for comprehensive RAG retrieval
            enriched = {
                "text": chunk.text,
                "start_index": chunk.start_index,
                "end_index": chunk.end_index,
                "token_count": chunk.token_count,
                "metadata": complete_metadata.copy()  # Include ALL metadata fields
            }
            enriched_chunks.append(enriched)
        return enriched_chunks
    except Exception as e:
        logger.error(f"Chunking failed: {e}")
        return []

def _make_thumb_b64(image_bytes: bytes, thumb_px: int = 256) -> str:
    """
    Returns a data-URI base64 thumbnail. Uses Pillow if available; otherwise returns
    a JPEG-compressed preview from the original bytes (best-effort).
    """
    try:
        if PIL_AVAILABLE:
            with Image.open(BytesIO(image_bytes)) as im:
                im = im.convert("RGB")
                im.thumbnail((thumb_px, thumb_px), Image.LANCZOS)
                out = BytesIO()
                im.save(out, format="JPEG", quality=80, optimize=True)
                b64 = base64.b64encode(out.getvalue()).decode("utf-8")
                return f"data:image/jpeg;base64,{b64}"
        # Fallback: just base64 the original (can be large)
        b64 = base64.b64encode(image_bytes).decode("utf-8")
        return f"data:image/{'jpeg'};base64,{b64}"
    except Exception:
        b64 = base64.b64encode(image_bytes).decode("utf-8")
        return f"data:image/{'jpeg'};base64,{b64}"

def build_pdf_image_manifest(pdf_path: str, guid: str,
                             min_area: int = 12000,
                             max_images: int = 150,
                             thumb_px: int = 256) -> list[dict]:
    """
    Scans PDF images and returns a manifest with small thumbnails and metadata.
    Does not persist full images to disk; we re-extract for selected items later.
    """
    manifest = []
    try:
        doc = fitz.open(pdf_path)
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            for img in page.get_images(full=True):
                xref = img[0]
                width, height = img[2], img[3]
                area = (width or 0) * (height or 0)
                if area < min_area:
                    continue
                try:
                    base_image = doc.extract_image(xref)
                except Exception:
                    continue
                image_bytes = base_image.get("image", b"")
                ext = base_image.get("ext", "png")
                size_bytes = len(image_bytes)
                thumb_b64 = _make_thumb_b64(image_bytes, thumb_px=thumb_px)
                entry_id = f"{guid}:{page_idx+1}:{xref}"
                manifest.append({
                    "id": entry_id,
                    "page": page_idx + 1,
                    "xref": xref,
                    "width": width,
                    "height": height,
                    "size_bytes": size_bytes,
                    "ext": ext,
                    "thumb_b64": thumb_b64
                })
                if len(manifest) >= max_images:
                    break
            if len(manifest) >= max_images:
                break
        doc.close()
    except Exception as e:
        logger.error(f"build_pdf_image_manifest error for {pdf_path}: {e}", exc_info=True)
    return manifest

def _selection_state_path(token: str) -> str:
    return os.path.join(JSON_CACHE_DIR, f"image_selection_{token}.json")

def save_selection_state(token: str, state: dict) -> None:
    try:
        with open(_selection_state_path(token), "w", encoding="utf-8") as f:
            json.dump(state, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.warning(f"Failed to save selection state: {e}")

def load_selection_state(token: str) -> dict | None:
    try:
        with open(_selection_state_path(token), "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None

def describe_selected_images_with_openai(pdf_path: str,
                                         selection_state: dict,
                                         selected_ids: list[str],
                                         model: str = "gpt-4o-mini",
                                         per_image_max_tokens: int = 250,
                                         sleep_sec: float = 0.0) -> list[dict]:
    """
    Calls OpenAI Vision only for selected ids.
    Returns list of {id, page, description}.
    """
    try:
        client = OpenAI()
    except Exception as e:
        logger.error(f"OpenAI client failed. Is OPENAI_API_KEY set? {e}")
        return []

    id_to_info = {e["id"]: e for e in selection_state.get("manifest", [])}
    results = []
    try:
        doc = fitz.open(pdf_path)
        for sel_id in selected_ids:
            info = id_to_info.get(sel_id)
            if not info:
                continue
            page = info.get("page")
            xref = info.get("xref")
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image.get("image", b"")
                b64 = base64.b64encode(image_bytes).decode("utf-8")
                prompt = (
                    "You are analyzing an image extracted from a PDF document. "
                    "Provide a precise, 1–3 sentence description tailored for retrieval. "
                    "If it is a chart/graph, state title (if visible), axes/units, key trend(s), and any standout values. "
                    "If it is a table/diagram/photo/illustration, summarize what it conveys. "
                    "Do not speculate beyond visible content."
                )
                resp = client.chat.completions.create(
                    model=model,
                    messages=[{
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}}
                        ]
                    }],
                    temperature=0.2,
                    max_tokens=per_image_max_tokens
                )
                desc = (resp.choices[0].message.content or "").strip()
                results.append({"id": sel_id, "page": page, "description": desc})
                if sleep_sec:
                    time.sleep(sleep_sec)
            except Exception as e:
                logger.warning(f"Vision call failed for image id={sel_id}: {e}")
        doc.close()
    except Exception as e:
        logger.error(f"Failed to open PDF for vision extraction: {e}", exc_info=True)
    return results

def append_image_descriptions_to_markdown(markdown_text: str, image_descs: list[dict]) -> str:
    if not image_descs:
        return markdown_text
    lines = []
    lines.append(markdown_text)
    lines.append("\n\n---\n### Image descriptions\n")
    # Sort by page then id for stable output
    for item in sorted(image_descs, key=lambda x: (x.get("page", 0), x.get("id", ""))):
        page = item.get("page")
        desc = item.get("description", "").strip()
        if desc:
            lines.append(f"- Page {page}: {desc}")
    return "\n".join(lines)

# --- API Endpoints ---


@app.route('/debug-paths', methods=['GET'])
def debug_paths():
    return jsonify({
        "BASE_DIR": BASE_DIR,
        "BASE_DIR_exists": os.path.exists(BASE_DIR),
        "BASE_DIR_absolute": os.path.abspath(BASE_DIR),
        "files_in_base": os.listdir(BASE_DIR) if os.path.exists(BASE_DIR) else "Directory not found"
    })

@app.route('/debug-file-search', methods=['GET'])
def debug_file_search():
    debug_info = {
        "BASE_DIR": BASE_DIR,
        "BASE_DIR_exists": os.path.exists(BASE_DIR),
        "SUPPORTED_EXTENSIONS": SUPPORTED_EXTENSIONS,
        "search_results": {}
    }
    
    for ext in SUPPORTED_EXTENSIONS:
        pattern = os.path.join(BASE_DIR, "**", f"*{ext}")
        files = glob.glob(pattern, recursive=True)
        debug_info["search_results"][ext] = {
            "pattern": pattern,
            "files_found": files
        }
    
    # Also try manual directory walk
    manual_files = []
    if os.path.exists(BASE_DIR):
        for root, dirs, files in os.walk(BASE_DIR):
            for file in files:
                if any(file.lower().endswith(ext.lower()) for ext in SUPPORTED_EXTENSIONS):
                    manual_files.append(os.path.join(root, file))
    
    debug_info["manual_walk_results"] = manual_files
    return jsonify(debug_info)

@app.route('/ensure-guids', methods=['POST'])
def ensure_guids():
    """Scans all supported files and ensures that any file missing a GUID gets one."""
    logger.info("Received request for /ensure-guids. Scanning all files.")
    try:
        all_files = find_all_supported_files()
        files_updated = 0
        
        for file_info in all_files:
            path = file_info['full_path']
            if path.lower().endswith('.docx'):
                if extract_docx_guid(path):
                    continue  # DOCX GUIDs are read-only
            elif path.lower().endswith('.pdf'):
                existing_guid = read_pdf_guid(path)
                if not existing_guid:
                    get_or_create_pdf_guid(path)
                    files_updated += 1

        message = f"Scan complete. Updated {files_updated} PDF file(s) with a new GUID."
        logger.info(message)
        return jsonify({"success": True, "message": message, "files_updated": files_updated})

    except Exception as e:
        logger.error(f"An error occurred in /ensure-guids: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/list-files', methods=['GET'])
def list_files():
    try:
        raw_file_list = find_all_supported_files()
        enriched_files = []

        for file_info in raw_file_list:
            path = file_info.get('full_path')
            if not path:
                logger.warning("Skipping file with no full_path in file_info")
                continue

            try:
                # Determine GUID/file_id and basic provenance without mutating files
                ext = os.path.splitext(path)[1].lower()
                guid = None
                guid_source = None
                date_reviewed = None
                date_next = None

                if ext == '.docx':
                    try:
                        guid = extract_docx_guid(path)
                        guid_source = 'docx_settings' if guid else 'filename_fallback'
                    except Exception as e:
                        logger.warning(f"Failed to read DOCX GUID foKr {path}: {e}")
                    try:
                        date_reviewed, date_next = extract_custom_metadata(path)
                    except Exception:
                        date_reviewed, date_next = None, None

                elif ext == '.pdf':
                    try:
                        guid = read_pdf_guid(path)
                        guid_source = 'pdf_metadata' if guid else 'filename_fallback'
                    except Exception as e:
                        logger.warning(f"Failed to read PDF GUID for {path}: {e}")

                # Fallback id when no GUID present
                if not guid:
                    filename_without_ext = os.path.splitext(os.path.basename(path))[0]
                    guid = f"fallback_{filename_without_ext.replace(' ', '_').lower()}"
                    if not guid_source:
                        guid_source = 'filename_fallback'

                extra = {
                    'file_id': guid,
                    'has_guid': bool(guid) and not str(guid).startswith('fallback_'),
                    'guid_source': guid_source,
                    'DateReviewed': date_reviewed,
                    'DateNext': date_next
                }

                # Build the full metadata dict using existing helper
                complete_meta = make_complete_metadata(path, extra)

                enriched_files.append(complete_meta)

            except Exception as inner_e:
                logger.error(f"Error processing file {path}: {inner_e}", exc_info=True)
                # Fallback minimal entry so list remains stable
                enriched_files.append({
                    'filename': os.path.basename(path) if path else None,
                    'full_path': path,
                    'file_id': None,
                    'lastModified': None,
                    'createdDate': None
                })

        return jsonify({"files": enriched_files, "count": len(enriched_files)})

    except Exception as e:
        logger.error(f"Error in /list-files: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500

@app.route('/docling/convert-file', methods=['POST'])
def docling_convert_file():
    """
    Convert a single file (DOCX or PDF) using its fileID as the unique identifier.
    Returns 200 OK with error flag on failure, allowing n8n to skip and continue.
    """
    file_id = None
    file_path = None
    try:
        # ===== STEP 1: Parse request =====
        try:
            data = request.get_json(force=True)
            if not data or 'file_id' not in data:
                return jsonify({"error": "Missing 'file_id' in JSON request body"}), 400
            file_id = data['file_id']
        except Exception as e:
            return jsonify({"error": f"Invalid request: {str(e)}"}), 400
        
        # ===== STEP 2: Resolve fileID to file_path (with auto-register fallback) =====
        try:
            file_path = resolve_file_path(file_id)
            
            # If not in registry, try to find file by ID and auto-register
            if not file_path:
                # Look for a file that matches this ID (could be from external source)
                # This supports calling convert-file with an ID even if list-files wasn't called first
                all_files = find_all_supported_files()
                for file_info in all_files:
                    candidate_path = file_info['full_path']
                    # Get the ID that would be assigned to this file
                    candidate_id = get_file_id_or_register(candidate_path)
                    if candidate_id == file_id:
                        file_path = candidate_path
                        break
                
                if not file_path:
                    return create_error_response(file_id, f"FileID not found and could not locate matching file: {file_id}", "FILE_ID_NOT_FOUND")
        except Exception as e:
            return create_error_response(file_id, f"Failed to resolve fileID: {str(e)}", "FILE_ID_RESOLUTION_ERROR")
        
        # ===== STEP 3: Validate file path =====
        try:
            if not os.path.exists(file_path):
                return create_error_response(file_id, f"File not found: {file_path}", "FILE_NOT_FOUND")
            if not os.path.abspath(file_path).startswith(os.path.abspath(BASE_DIR)):
                return create_error_response(file_id, "File path is outside allowed directory", "SECURITY_ERROR")
        except Exception as e:
            return create_error_response(file_id, f"Failed to validate file path: {str(e)}", "PATH_VALIDATION_ERROR")
        
        # ===== STEP 4: Extract extension =====
        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext not in SUPPORTED_EXTENSIONS:
                return create_error_response(file_id, f"Unsupported file type: {ext}", "UNSUPPORTED_FORMAT")
        except Exception as e:
            return create_error_response(file_id, f"Failed to determine file type: {str(e)}", "FILE_TYPE_ERROR")

        # ===== STEP 5: File timestamps (UTC) =====
        try:
            stat = os.stat(file_path)
            last_modified_iso = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat()
            created_date_iso = datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc).isoformat()
        except Exception as e:
            logger.warning(f"Failed to stat file for timestamps: {e}")
            last_modified_iso = None
            created_date_iso = None

        # [Existing Params]
        force_refresh = data.get("force_refresh", False)
        do_chunking = data.get("do_chunking", False)
        chunking_options = data.get("chunking_options", {})
        image_mode = data.get("image_mode", "ask")
        selection_token = data.get("selection_token")
        selected_image_ids = data.get("selected_image_ids", []) or []

        # ===== STEP 6: Count images (PDF & DOCX) =====
        # Fast pre-check to see what we are dealing with
        raw_total_images = 0
        try:
            if ext == '.pdf':
                with fitz.open(file_path) as count_doc:
                    for page in count_doc:
                        raw_total_images += len(page.get_images(full=True))
            
            elif ext == '.docx':
                try:
                    with zipfile.ZipFile(file_path, 'r') as z:
                        media_files = [n for n in z.namelist() if n.startswith('word/media/')]
                        raw_total_images = len(media_files)
                except Exception as e:
                    logger.warning(f"Could not count images in DOCX {file_path}: {e}")
                    
        except Exception as e:
            logger.warning(f"Image counting failed for {file_path}: {e}")

        # ===== STEP 7: SURVEY MODE (image_mode="ask") =====
        # If the workflow asks us to survey, we return metadata immediately.
        # We do NOT run Docling here. We wait for user confirmation.
        if image_mode == "ask":
            logger.info(f"Surveying file (Ask Mode): {file_id} - Images found: {raw_total_images}")
            
            
            # 1. Prepare Image Manifest (if images exist)
            image_manifest = []
            selection_token = None
            
            if ext == ".pdf" and raw_total_images > 0:
                try:
                    image_manifest = build_pdf_image_manifest(file_path, file_id)
                    if image_manifest:
                        token_seed = f"{file_id}:{file_path}:{time.time()}:{len(image_manifest)}"
                        selection_token = hashlib.sha256(token_seed.encode("utf-8")).hexdigest()[:32]
                        
                        # Save state so we can retrieve these specific images later
                        save_selection_state(selection_token, {
                            "file_path": file_path,
                            "file_id": file_id,
                            "manifest": image_manifest,
                            "created_at": datetime.now(timezone.utc).isoformat()
                        })
                except Exception as e:
                    logger.warning(f"Failed to build image manifest: {e}")

            # 2. Return 'Pending Approval' JSON
            # This stops the script here. The UI will pick this up.
            return jsonify({
                "success": True,
                "status": "pending_approval",
                "skipped": True,
                "file_path": file_path,
                "filename": os.path.basename(file_path),
                "lastModified": last_modified_iso,
                "createdDate": created_date_iso,
                "file_id": file_id,
                "image_count_total": raw_total_images,
                
                # Image Selection Data
                "needs_image_selection": (len(image_manifest) > 0),
                "image_manifest": image_manifest,
                "selection_token": selection_token
            })

        # The code below only runs if image_mode is "selected" or "force"
        # (i.e. AFTER the user has clicked 'Proceed' in the UI).
        # ==================================================================
        # [SLOW PATH] STANDARD CONVERSION
        # ==================================================================
        
        # FIX: Define use_docling here. 
        # Since you want Docling for all files, we set this to True.
        # This ensures the script skips "[Path A: Native DOCX]" below and goes straight to Docling.
        use_docling = True

        # [Path A: Native DOCX]
        if ext == '.docx' and not use_docling:
            try:
                try:
                    content = docx_to_json(file_path)
                except Exception as e:
                    return create_error_response(file_id, f"Failed to convert DOCX: {str(e)}", "DOCX_CONVERSION_ERROR", {"step": "docx_to_json"})
                
                if "error" in content:
                    return create_error_response(file_id, content.get("error", "Unknown error in DOCX conversion"), "DOCX_PARSING_ERROR", {"step": "docx_parsing"})
                
                try:
                    date_reviewed, date_next = extract_custom_metadata(file_path)
                except Exception as e:
                    logger.warning(f"Failed to extract custom metadata from DOCX: {e}")
                    date_reviewed, date_next = None, None
                
                try:
                    full_text = "\n\n".join(content.get("paragraphs", []))
                except Exception as e:
                    return create_error_response(file_id, f"Failed to extract text from DOCX: {str(e)}", "TEXT_EXTRACTION_ERROR")
                
                try:
                    ai_title, ai_description = generate_ai_title_description(full_text, os.path.basename(file_path))
                except Exception as e:
                    logger.warning(f"Failed to generate AI title/description: {e}")
                    ai_title, ai_description = None, None
                
                extra_data = {
                    "file_id": file_id,
                    "paragraph_count": len(content.get("paragraphs", [])),
                    "table_count": len(content.get("tables", [])),
                    "word_count": len(full_text.split()),
                    "character_count": len(full_text),
                    "has_guid": True,
                    "guid_source": "file_id",
                    "DateReviewed": date_reviewed,
                    "DateNext": date_next,
                    "ai_title": ai_title,
                    "ai_description": ai_description,
                    "image_count_total": 0,
                    "image_count_selected": 0,
                    "image_count_analyzed": 0,
                }
                complete_metadata = make_complete_metadata(file_path, extra_data)
                chunks = []
                if do_chunking:
                    try:
                        chunks = perform_chunking(full_text, complete_metadata, chunking_options)
                    except Exception as e:
                        logger.warning(f"Failed to chunk DOCX: {e}")
                        chunks = []

                return jsonify({
                    "success": True,
                    "used_converter": "native_docx",
                    "file_path": file_path,
                    "filename": content.get("filename"),
                    "full_text": full_text,
                    "structured_content": {
                        "paragraphs": content.get("paragraphs"),
                        "tables": content.get("tables"),
                        "metadata": content.get("metadata")
                    },
                    "metadata": complete_metadata,
                    "chunks": chunks,
                    "lastModified": last_modified_iso,
                    "createdDate": created_date_iso
                })
            except Exception as e:
                return create_error_response(file_id, f"DOCX conversion failed: {str(e)}", "DOCX_CONVERSION_FAILED", {"exception_type": type(e).__name__})

        # [Path B: Heavy Docling Conversion]
        try:
            logger.info("Running Docling converter (Heavy Process)...")
            try:
                conversion_result = docling_converter.convert(file_path)
            except Exception as e:
                return create_error_response(file_id, f"Docling conversion failed: {str(e)}", "DOCLING_CONVERSION_ERROR", {"exception_type": type(e).__name__})
            
            try:
                document = conversion_result.document
                document_dict = document.export_to_dict()
                full_text_md = document.export_to_markdown()
            except Exception as e:
                return create_error_response(file_id, f"Failed to export document: {str(e)}", "DOCUMENT_EXPORT_ERROR", {"exception_type": type(e).__name__})

            extra_meta = {}
            if ext == ".docx":
                try:
                    date_reviewed, date_next = extract_custom_metadata(file_path)
                    extra_meta.update({"DateReviewed": date_reviewed, "DateNext": date_next})
                except Exception as e:
                    logger.warning(f"Failed to extract custom metadata: {e}")

            # Handle Image Mode = SELECTED
            image_descs = []
            selection_state = None
            
            if ext == ".pdf" and image_mode == "selected":
                try:
                    if not selection_token:
                        return create_error_response(file_id, "selection_token is required when image_mode='selected'", "MISSING_TOKEN")
                    
                    selection_state = load_selection_state(selection_token)
                    if not selection_state or selection_state.get("file_path") != file_path:
                        return create_error_response(file_id, "Invalid or expired selection_token", "INVALID_TOKEN")

                    try:
                        allowed_keys = {"pdf_path", "selection_state", "selected_ids", "model", "per_image_max_tokens", "sleep_sec"}
                        describe_kwargs = {
                            "pdf_path": file_path,
                            "selection_state": selection_state,
                            "selected_ids": selected_image_ids,
                            "model": data.get("vision_model", "gpt-4o-mini"),
                            "per_image_max_tokens": int(data.get("vision_max_tokens", 250)),
                            "sleep_sec": float(data.get("vision_sleep_sec", 0.0)),
                        }
                        describe_kwargs = {k: v for k, v in describe_kwargs.items() if k in allowed_keys}
                    
                        # Analyze images
                        image_descs = describe_selected_images_with_openai(**describe_kwargs)
                    except Exception as e:
                        logger.warning(f"Failed to analyze selected images: {e}")
                        image_descs = []
                    
                    try:
                        # Append descriptions to Markdown
                        full_text_md = append_image_descriptions_to_markdown(full_text_md, image_descs)
                    except Exception as e:
                        logger.warning(f"Failed to append image descriptions: {e}")
                except Exception as e:
                    logger.warning(f"Image analysis phase failed (non-fatal): {e}")

            # Final Metadata & Response Construction
            try:
                ai_title, ai_description = generate_ai_title_description(full_text_md, Path(file_path).name)
            except Exception as e:
                logger.warning(f"Failed to generate AI title/description: {e}")
                ai_title, ai_description = None, None

            image_count_selected = len(selected_image_ids) if selected_image_ids else 0
            image_count_analyzed = len(image_descs) if image_descs else 0

            extra_data = {
                "file_id": file_id,
                "paragraph_count": full_text_md.count('\n\n') + 1 if full_text_md else 0,
                "table_count": len(document_dict.get("tables", [])) if document_dict else 0,
                "word_count": len(full_text_md.split()) if full_text_md else 0,
                "character_count": len(full_text_md) if full_text_md else 0,
                "has_guid": True,
                "guid_source": "file_id",
                "ai_title": ai_title,
                "ai_description": ai_description,
                "image_count_total": raw_total_images,
                "image_count_selected": image_count_selected,
                "image_count_analyzed": image_count_analyzed,
                **extra_meta,
            }

            origin_meta_raw = document_dict.get("origin", {}) if document_dict else {}
            if isinstance(origin_meta_raw, dict):
                extra_data.update(origin_meta_raw)
            elif origin_meta_raw:
                extra_data["origin"] = str(origin_meta_raw)

            complete_metadata = make_complete_metadata(file_path, extra_data)
            
            chunks = []
            if do_chunking:
                try:
                    chunks = perform_chunking(full_text_md, complete_metadata, chunking_options)
                except Exception as e:
                    logger.warning(f"Failed to chunk document: {e}")
                    chunks = []

            resp = {
                "success": True,
                "used_converter": "docling",
                "file_id": file_id,
                "file_path": file_path,
                "filename": os.path.basename(file_path),
                "conversion_status": str(conversion_result.status) if conversion_result else "unknown",
                "document": document_dict,
                "full_text": full_text_md,
                "metadata": complete_metadata,
                "chunks": chunks
            }
            # Include timestamps at top-level for easy access
            resp["lastModified"] = last_modified_iso
            resp["createdDate"] = created_date_iso
            if selection_state:
                resp["image_manifest"] = selection_state.get("manifest", [])
            if image_descs:
                resp["image_descriptions"] = image_descs

            return jsonify(resp)

        except Exception as e:
            return create_error_response(file_id, f"Docling conversion failed: {str(e)}", "DOCLING_CONVERSION_FAILED", {"exception_type": type(e).__name__})

    except Exception as e:
        # Catch-all for any unexpected errors
        logger.error(f"Unexpected error in docling_convert_file: {e}", exc_info=True)
        return create_error_response(file_id or "unknown", f"Unexpected server error: {str(e)}", "UNEXPECTED_ERROR", {"exception_type": type(e).__name__})

@app.route('/file-id/resolve', methods=['GET'])
def resolve_file_endpoint():
    """Resolve a fileID to its file path."""
    try:
        file_id = request.args.get('file_id')
        if not file_id:
            return jsonify({"error": "Missing 'file_id' parameter", "success": False}), 400
        
        file_path = resolve_file_path(file_id)
        
        if not file_path:
            return jsonify({"error": f"FileID not found: {file_id}", "success": False}), 404
        
        return jsonify({
            "success": True,
            "file_id": file_id,
            "file_path": file_path,
            "filename": os.path.basename(file_path)
        }), 200
    except Exception as e:
        logger.error(f"Error in /file-id/resolve: {e}", exc_info=True)
        return jsonify({"error": str(e), "success": False}), 500

@app.route('/clear-cache', methods=['POST'])
def clear_cache():
    try:
        if not os.path.exists(JSON_CACHE_DIR):
            return jsonify({"success": True, "message": "Cache directory doesn't exist, nothing to clear", "files_deleted": 0})
        cache_files = glob.glob(os.path.join(JSON_CACHE_DIR, "*.json"))
        files_deleted = 0
        for cache_file in cache_files:
            try:
                os.remove(cache_file)
                files_deleted += 1
            except Exception as e:
                logger.warning(f"Could not delete {cache_file}: {e}")
        return jsonify({"success": True, "message": f"Cache cleared successfully", "files_deleted": files_deleted, "cache_directory": JSON_CACHE_DIR})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/health', methods=['GET'])
def health():
    return jsonify({"status": "ok", "base_directory": BASE_DIR})

# Docling configuration
pdf_opts = PdfPipelineOptions(
    do_ocr=True,
    do_table_structure=True,
    table_structure_options=TableStructureOptions(
        mode=TableFormerMode.ACCURATE
    )
)

docling_converter = DocumentConverter(
    format_options={
        InputFormat.PDF:  PdfFormatOption(pipeline_options=pdf_opts),
        InputFormat.DOCX: WordFormatOption(),
    }
)

    
@app.route('/image-selection/get', methods=['GET'])
def image_selection_get():
    token = request.args.get('token')
    if not token:
        return jsonify({"error": "missing token"}), 400
    state = load_selection_state(token)
    if not state:
        return jsonify({"error": "invalid or expired token"}), 404
    return jsonify({
        "file_path": state.get("file_path"),
        "guid": state.get("guid"),
        "image_manifest": state.get("manifest", []),
        "created_at": state.get("created_at")
    })

@app.route('/docling/convert-all', methods=['POST'])
def docling_convert_all():
    """
    Batch-convert all supported files under BASE_DIR using the Docling path.
    Gracefully skips corrupt/unreadable files and continues with the next.
    Returns detailed results for each file processed.
    """
    results = {
        "total_files": 0,
        "successful": 0,
        "failed": 0,
        "skipped": 0,
        "file_results": []
    }
    
    try:
        files = find_all_supported_files()
        results["total_files"] = len(files)
        
        for file_info in files:
            path = file_info.get("full_path")
            if not path:
                logger.warning("Skipping file with no path in file_info")
                continue
            
            file_result = {
                "file_path": path,
                "filename": os.path.basename(path),
                "status": "pending",
                "success": False,
                "error": None,
                "error_type": None
            }
            
            try:
                # Validate file existence and readability
                if not os.path.exists(path):
                    file_result["status"] = "skipped"
                    file_result["error"] = "File not found"
                    file_result["error_type"] = "FILE_NOT_FOUND"
                    results["skipped"] += 1
                    results["file_results"].append(file_result)
                    logger.warning(f"Skipping: file not found: {path}")
                    continue
                
                # Check file is readable
                if not os.access(path, os.R_OK):
                    file_result["status"] = "skipped"
                    file_result["error"] = "File is not readable"
                    file_result["error_type"] = "PERMISSION_ERROR"
                    results["skipped"] += 1
                    results["file_results"].append(file_result)
                    logger.warning(f"Skipping: permission denied: {path}")
                    continue
                
                # Attempt conversion
                try:
                    logger.info(f"Converting: {path}")
                    conversion_result = docling_converter.convert(path)
                    
                    file_result["status"] = "completed"
                    file_result["success"] = True
                    file_result["conversion_status"] = str(conversion_result.status)
                    results["successful"] += 1
                    logger.info(f"Successfully converted: {path}")
                    
                except Exception as convert_error:
                    file_result["status"] = "failed"
                    file_result["success"] = False
                    file_result["error"] = str(convert_error)
                    file_result["error_type"] = type(convert_error).__name__
                    results["failed"] += 1
                    logger.error(f"Conversion failed for {path}: {convert_error}")
            
            except Exception as file_error:
                file_result["status"] = "error"
                file_result["success"] = False
                file_result["error"] = str(file_error)
                file_result["error_type"] = type(file_error).__name__
                results["failed"] += 1
                logger.error(f"Error processing {path}: {file_error}", exc_info=True)
            
            finally:
                results["file_results"].append(file_result)
        
        # Summary response
        return jsonify({
            "success": True,
            "message": f"Batch conversion complete: {results['successful']} succeeded, {results['failed']} failed, {results['skipped']} skipped",
            "results": results
        }), 200
    
    except Exception as e:
        logger.error(f"Error in batch conversion: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": f"Batch conversion error: {str(e)}",
            "error_type": type(e).__name__,
            "results": results
        }), 200  # Still return 200 even on batch-level error, with partial results

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5001, debug=True, use_reloader=False)